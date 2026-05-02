"""
Image-to-Word Converter - Phase 1 MVP
FAST-NUCES BSAI Project
Converts scanned/photographed documents to formatted .docx files
using Tesseract OCR with formatting detection.
"""

import os
import sys
import re
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from PIL import Image, ImageTk
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
import pytesseract
import cv2
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import threading


# ─── OCR & Preprocessing ────────────────────────────────────────────────────

def preprocess_image(image_path: str) -> np.ndarray:
    """
    Preprocess image for better OCR accuracy:
    - Convert to grayscale
    - Denoise
    - Adaptive thresholding (handles uneven lighting)
    - Deskew
    """
    img = cv2.imread(image_path)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    # Denoise
    denoised = cv2.fastNlMeansDenoising(gray, h=10)

    # Adaptive threshold for uneven lighting
    thresh = cv2.adaptiveThreshold(
        denoised, 255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY, 31, 10
    )

    # Deskew
    coords = np.column_stack(np.where(thresh < 128))
    if len(coords) > 0:
        angle = cv2.minAreaRect(coords)[-1]
        if angle < -45:
            angle = -(90 + angle)
        else:
            angle = -angle
        if abs(angle) < 15:  # Only correct small skews
            (h, w) = thresh.shape
            center = (w // 2, h // 2)
            M = cv2.getRotationMatrix2D(center, angle, 1.0)
            thresh = cv2.warpAffine(
                thresh, M, (w, h),
                flags=cv2.INTER_CUBIC,
                borderMode=cv2.BORDER_REPLICATE
            )

    return thresh


def extract_with_layout(image_path: str) -> list[dict]:
    """
    Run Tesseract with full layout data (hOCR-style via image_to_data).
    Returns list of word dicts with bounding boxes, confidence, text, etc.
    """
    preprocessed = preprocess_image(image_path)
    pil_img = Image.fromarray(preprocessed)

    data = pytesseract.image_to_data(
        pil_img,
        output_type=pytesseract.Output.DICT,
        config='--psm 6'  # Assume uniform block of text
    )
    return data


def detect_bold(word: str, width: int, height: int, conf: float) -> bool:
    """
    Heuristic bold detection:
    - ALL CAPS words are treated as bold/heading
    - Words with high aspect ratio (wide relative to height) suggest bold font weight
    """
    if not word.strip():
        return False
    if word.isupper() and len(word) > 1:
        return True
    if height > 0 and width / height > 0.75:
        return True
    return False


def detect_italic(word: str) -> bool:
    """
    Heuristic italic detection:
    - Words surrounded by *asterisks* or _underscores_
    - Words ending with common italic markers
    """
    return bool(re.match(r'^\*[^*]+\*$', word) or re.match(r'^_[^_]+_$', word))


def group_into_lines(data: dict) -> list[list[dict]]:
    """
    Group Tesseract word data into lines using line_num and block_num.
    Returns list of lines, each line being a list of word dicts.
    """
    lines = {}
    n = len(data['text'])
    for i in range(n):
        if int(data['conf'][i]) < 20:
            continue
        word = data['text'][i].strip()
        if not word:
            continue

        block = data['block_num'][i]
        line = data['line_num'][i]
        key = (block, line)

        entry = {
            'text': word,
            'left': data['left'][i],
            'top': data['top'][i],
            'width': data['width'][i],
            'height': data['height'][i],
            'conf': float(data['conf'][i]),
        }
        entry['bold'] = detect_bold(
            word, entry['width'], entry['height'], entry['conf']
        )
        entry['italic'] = detect_italic(word)

        if key not in lines:
            lines[key] = []
        lines[key].append(entry)

    # Sort lines by vertical position
    sorted_keys = sorted(lines.keys(), key=lambda k: (
        min(w['top'] for w in lines[k])
    ))
    return [lines[k] for k in sorted_keys]


def detect_alignment(line_words: list[dict], page_width: int) -> str:
    """
    Detect paragraph alignment:
    - If left edge starts near center → centered
    - If text starts far from left margin → right aligned
    - Otherwise → left aligned
    """
    if not line_words:
        return 'LEFT'

    leftmost = min(w['left'] for w in line_words)
    rightmost = max(w['left'] + w['width'] for w in line_words)
    center_pos = (leftmost + rightmost) / 2

    if page_width > 0:
        if abs(center_pos - page_width / 2) < page_width * 0.08:
            return 'CENTER'
        if leftmost > page_width * 0.55:
            return 'RIGHT'
    return 'LEFT'


def detect_heading(line_words: list[dict], avg_height: float) -> bool:
    """A line is a heading if its average word height is significantly larger."""
    if not line_words:
        return False
    line_avg = sum(w['height'] for w in line_words) / len(line_words)
    return line_avg > avg_height * 1.4


# ─── Word Document Generation ────────────────────────────────────────────────

ALIGN_MAP = {
    'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
    'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
    'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
    'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def build_docx(lines: list[list[dict]], page_width: int, output_path: str):
    """
    Build a formatted .docx from grouped line data.
    Applies bold, italic, alignment, and heading styles.
    """
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Compute average word height for heading detection
    all_heights = [w['height'] for line in lines for w in line if w['height'] > 0]
    avg_height = sum(all_heights) / len(all_heights) if all_heights else 20

    for line_words in lines:
        if not line_words:
            continue

        line_text = ' '.join(w['text'] for w in line_words)
        alignment = detect_alignment(line_words, page_width)
        is_heading = detect_heading(line_words, avg_height)

        if is_heading:
            para = doc.add_heading(level=1)
        else:
            para = doc.add_paragraph()

        para.alignment = ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

        # Add words with formatting
        for word in line_words:
            run = para.add_run(word['text'] + ' ')
            run.bold = word['bold'] or is_heading
            run.italic = word['italic']
            if not is_heading:
                run.font.size = Pt(12)

    doc.save(output_path)


# ─── Full Pipeline ────────────────────────────────────────────────────────────

def convert_image_to_docx(image_path: str, output_path: str, progress_cb=None) -> str:
    """
    Full pipeline: image → preprocess → OCR → format detect → .docx
    Returns output path on success.
    """
    def progress(msg):
        if progress_cb:
            progress_cb(msg)

    progress("🔍 Preprocessing image...")
    preprocessed = preprocess_image(image_path)

    # Get page width from original image
    orig = cv2.imread(image_path)
    page_width = orig.shape[1] if orig is not None else 800

    progress("📖 Running Tesseract OCR...")
    pil_img = Image.fromarray(preprocessed)
    data = pytesseract.image_to_data(
        pil_img,
        output_type=pytesseract.Output.DICT,
        config='--psm 6 --oem 3'
    )

    progress("🔎 Detecting formatting (bold, italic, alignment)...")
    lines = group_into_lines(data)

    if not lines:
        raise ValueError("No text could be extracted from the image.")

    progress("📝 Generating Word document...")
    build_docx(lines, page_width, output_path)

    progress(f"✅ Done! Saved to: {output_path}")
    return output_path


# ─── GUI ─────────────────────────────────────────────────────────────────────

class ImageToWordApp(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title("Image → Word Converter  |  FAST-NUCES BSAI")
        self.geometry("820x700")
        self.resizable(True, True)
        self.configure(bg="#0F172A")

        self._image_path = None
        self._output_path = None
        self._thumb = None

        self._build_ui()

    # ── UI Construction ──────────────────────────────────────────────────────

    def _build_ui(self):
        # ── Header ──
        header = tk.Frame(self, bg="#1E293B", pady=14)
        header.pack(fill="x")

        tk.Label(
            header,
            text="⬛ IMAGE → WORD CONVERTER",
            font=("Courier New", 15, "bold"),
            fg="#38BDF8", bg="#1E293B"
        ).pack()
        tk.Label(
            header,
            text="FAST-NUCES  ·  BSAI  ·  Phase 1 MVP",
            font=("Courier New", 9),
            fg="#64748B", bg="#1E293B"
        ).pack()

        # ── Main body ──
        body = tk.Frame(self, bg="#0F172A", padx=24, pady=16)
        body.pack(fill="both", expand=True)

        # Left: Drop zone / preview
        left = tk.Frame(body, bg="#0F172A")
        left.pack(side="left", fill="both", expand=True)

        self._drop_frame = tk.Frame(
            left, bg="#1E293B", relief="flat",
            bd=0, padx=10, pady=10
        )
        self._drop_frame.pack(fill="both", expand=True, pady=(0, 10))

        self._preview_label = tk.Label(
            self._drop_frame,
            text="📂\n\nClick Browse to select\nan image file\n\nJPG  ·  PNG  ·  JPEG",
            font=("Courier New", 11),
            fg="#94A3B8", bg="#1E293B",
            justify="center"
        )
        self._preview_label.pack(expand=True)

        # Right: Controls
        right = tk.Frame(body, bg="#0F172A", width=240)
        right.pack(side="right", fill="y", padx=(16, 0))
        right.pack_propagate(False)

        self._make_section(right, "INPUT")
        self._btn_browse = self._make_button(
            right, "📁  Browse Image", self._browse_image, accent=True
        )

        self._file_label = tk.Label(
            right, text="No file selected",
            font=("Courier New", 8), fg="#64748B", bg="#0F172A",
            wraplength=220, justify="left"
        )
        self._file_label.pack(anchor="w", pady=(4, 12))

        self._make_section(right, "OUTPUT")
        self._btn_output = self._make_button(
            right, "📂  Choose Output", self._choose_output
        )
        self._output_label = tk.Label(
            right, text="Auto-named alongside input",
            font=("Courier New", 8), fg="#64748B", bg="#0F172A",
            wraplength=220, justify="left"
        )
        self._output_label.pack(anchor="w", pady=(4, 12))

        self._make_section(right, "CONVERT")
        self._btn_convert = self._make_button(
            right, "⚡  CONVERT NOW", self._start_convert, accent=True, big=True
        )
        self._btn_convert.config(state="disabled")

        # Progress / log
        self._make_section(right, "LOG")
        self._log_text = tk.Text(
            right, height=7, bg="#1E293B", fg="#94A3B8",
            font=("Courier New", 8), relief="flat",
            state="disabled", wrap="word", bd=0
        )
        self._log_text.pack(fill="x")

        # Status bar
        self._status = tk.Label(
            self, text="Ready",
            font=("Courier New", 9), fg="#38BDF8",
            bg="#1E293B", anchor="w", padx=12, pady=6
        )
        self._status.pack(fill="x", side="bottom")

        # Progress bar
        style = ttk.Style()
        style.theme_use("clam")
        style.configure(
            "Blue.Horizontal.TProgressbar",
            troughcolor="#1E293B", background="#38BDF8",
            thickness=4
        )
        self._progress = ttk.Progressbar(
            self, style="Blue.Horizontal.TProgressbar",
            mode="indeterminate"
        )
        self._progress.pack(fill="x", side="bottom")

    def _make_section(self, parent, label):
        tk.Label(
            parent, text=f"── {label} ──",
            font=("Courier New", 8, "bold"),
            fg="#334155", bg="#0F172A"
        ).pack(anchor="w", pady=(10, 2))

    def _make_button(self, parent, text, cmd, accent=False, big=False):
        bg = "#0EA5E9" if accent else "#1E293B"
        fg = "#0F172A" if accent else "#94A3B8"
        size = 10 if big else 9
        btn = tk.Button(
            parent, text=text, command=cmd,
            font=("Courier New", size, "bold" if big else "normal"),
            bg=bg, fg=fg, activebackground="#38BDF8",
            activeforeground="#0F172A",
            relief="flat", pady=8, padx=6,
            cursor="hand2", bd=0
        )
        btn.pack(fill="x", pady=2)
        return btn

    # ── Callbacks ────────────────────────────────────────────────────────────

    def _browse_image(self):
        path = filedialog.askopenfilename(
            title="Select Image",
            filetypes=[("Image files", "*.jpg *.jpeg *.png"), ("All files", "*.*")]
        )
        if not path:
            return
        self._image_path = path
        self._file_label.config(text=os.path.basename(path))
        self._show_preview(path)

        # Auto-set output path
        base, _ = os.path.splitext(path)
        self._output_path = base + "_converted.docx"
        self._output_label.config(text=os.path.basename(self._output_path))

        self._btn_convert.config(state="normal")
        self._log("Image loaded: " + os.path.basename(path))

    def _choose_output(self):
        path = filedialog.asksaveasfilename(
            title="Save Word Document As",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")]
        )
        if path:
            self._output_path = path
            self._output_label.config(text=os.path.basename(path))

    def _show_preview(self, path):
        try:
            img = Image.open(path)
            img.thumbnail((400, 380))
            self._thumb = ImageTk.PhotoImage(img)
            self._preview_label.config(image=self._thumb, text="")
        except Exception as e:
            self._preview_label.config(text=f"Preview error:\n{e}", image="")

    def _start_convert(self):
        if not self._image_path:
            messagebox.showwarning("No Image", "Please select an image first.")
            return

        self._btn_convert.config(state="disabled")
        self._btn_browse.config(state="disabled")
        self._progress.start(12)
        self._status.config(text="Converting...")
        self._log("─" * 32)

        thread = threading.Thread(target=self._run_convert, daemon=True)
        thread.start()

    def _run_convert(self):
        try:
            convert_image_to_docx(
                self._image_path,
                self._output_path,
                progress_cb=lambda msg: self.after(0, self._log, msg)
            )
            self.after(0, self._on_success)
        except Exception as e:
            self.after(0, self._on_error, str(e))

    def _on_success(self):
        self._progress.stop()
        self._status.config(text=f"✅ Saved: {self._output_path}")
        self._btn_convert.config(state="normal")
        self._btn_browse.config(state="normal")
        messagebox.showinfo(
            "Conversion Complete",
            f"Word document saved to:\n{self._output_path}"
        )

    def _on_error(self, msg):
        self._progress.stop()
        self._status.config(text="❌ Error occurred")
        self._btn_convert.config(state="normal")
        self._btn_browse.config(state="normal")
        self._log(f"ERROR: {msg}")
        messagebox.showerror("Conversion Failed", msg)

    def _log(self, msg):
        self._log_text.config(state="normal")
        self._log_text.insert("end", msg + "\n")
        self._log_text.see("end")
        self._log_text.config(state="disabled")


# ─── Entry Point ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    app = ImageToWordApp()
    app.mainloop()
