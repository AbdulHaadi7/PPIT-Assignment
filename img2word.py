"""
Image → Word Converter | Phase 1 MVP
FAST-NUCES BSAI Virtual Company Project
OCR : Surya 0.17.x  (OCRPredictor, local, zero quota)
LLM : Groq  (LLaMA 3.3 70B, free tier, formats text)
UI  : Gradio → HuggingFace Spaces
"""

import os
import re
import io
import hashlib
import functools

import gradio as gr
from PIL import Image, ImageEnhance
from groq import Groq
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Lazy-load Surya 0.4.x models (downloaded from HuggingFace, cached) ───────
@functools.lru_cache(maxsize=1)
def get_surya_models():
    from surya.model.detection.segformer import load_model as load_det_model, load_processor as load_det_processor
    from surya.model.recognition.model import load_model as load_rec_model
    from surya.model.recognition.processor import load_processor as load_rec_processor
    det_processor = load_det_processor()
    det_model     = load_det_model()
    rec_model     = load_rec_model()
    rec_processor = load_rec_processor()
    return det_model, det_processor, rec_model, rec_processor


# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 — Surya OCR  (Vision Transformer, runs locally, no API quota)
# ══════════════════════════════════════════════════════════════════════════════

def run_surya_ocr(image: Image.Image) -> str:
    """
    Extract raw text using Surya 0.4.x run_ocr API.
    Returns plain text, one line per detected text line,
    in top-to-bottom reading order.
    """
    from surya.ocr import run_ocr
    image = ImageEnhance.Contrast(image).enhance(1.3)
    if image.mode != "RGB":
        image = image.convert("RGB")

    det_model, det_processor, rec_model, rec_processor = get_surya_models()
    results = run_ocr([image], [["en"]], det_model, det_processor, rec_model, rec_processor)

    lines = [line.text for line in results[0].text_lines if line.text.strip()]
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 — Groq LLM  (formats raw OCR text into markdown)
# ══════════════════════════════════════════════════════════════════════════════

GROQ_SYSTEM = """You are a document formatting assistant.
You receive raw OCR-extracted text from a scanned document and must reformat it.

Rules:
1. Use **text** for headings, titles, or bold content
2. Use *text* for italic content
3. Use "• " prefix for bullet points
4. Preserve numbered lists as-is
5. Separate paragraphs with a blank line
6. Fix obvious OCR errors (broken words, stray characters)
7. Preserve all content — do NOT summarize or remove text
8. Output ONLY the formatted text — no commentary, no preamble
"""

def format_with_groq(raw_text: str, api_key: str) -> str:
    client = Groq(api_key=api_key)
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": GROQ_SYSTEM},
            {"role": "user",   "content": f"Format this OCR text:\n\n{raw_text}"}
        ],
        temperature=0.2,
        max_tokens=4096,
    )
    return response.choices[0].message.content.strip()


# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 — Build .docx from formatted markdown text
# ══════════════════════════════════════════════════════════════════════════════

def parse_formatting(line: str):
    """Parse **bold**, *italic*, plain segments from a line."""
    segments = []
    for m in re.finditer(r'\*\*(.+?)\*\*|\*(.+?)\*|([^*]+)', line):
        if m.group(1):
            segments.append((m.group(1), True, False))
        elif m.group(2):
            segments.append((m.group(2), False, True))
        elif m.group(3) and m.group(3).strip():
            segments.append((m.group(3), False, False))
    return segments


def is_heading(line: str) -> bool:
    s = line.strip()
    return bool(re.match(r'^\*\*.+\*\*$', s)) or s.startswith('#')


def build_docx(formatted_text: str, font_size: int = 12) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(1)
    sec.top_margin  = sec.bottom_margin = Inches(1)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(font_size)

    for line in formatted_text.split('\n'):
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()
            continue

        hm = re.match(r'^(#{1,3})\s+(.+)', stripped)
        if hm:
            level = len(hm.group(1))
            run = doc.add_heading(level=level).add_run(
                hm.group(2).replace('**','').replace('*',''))
            run.bold = True
            continue

        if is_heading(stripped) and len(stripped.replace('**','')) < 70:
            run = doc.add_heading(level=1).add_run(
                stripped.replace('**','').replace('*',''))
            run.bold = True
            continue

        if stripped.startswith(('•', '-', '*  ')):
            para = doc.add_paragraph(style='List Bullet')
            content = re.sub(r'^[•\-\*]\s*', '', stripped)
            for text, bold, italic in parse_formatting(content):
                r = para.add_run(text)
                r.bold = bold; r.italic = italic; r.font.size = Pt(font_size)
            continue

        if re.match(r'^\d+[\.\)]\s', stripped):
            para = doc.add_paragraph(style='List Number')
            content = re.sub(r'^\d+[\.\)]\s*', '', stripped)
            for text, bold, italic in parse_formatting(content):
                r = para.add_run(text)
                r.bold = bold; r.italic = italic; r.font.size = Pt(font_size)
            continue

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for text, bold, italic in parse_formatting(stripped):
            r = para.add_run(text)
            r.bold = bold; r.italic = italic; r.font.size = Pt(font_size)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# PIPELINE  (OCR → LLM → docx)
# ══════════════════════════════════════════════════════════════════════════════

_cache: dict = {}

def get_image_hash(image: Image.Image) -> str:
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    return hashlib.md5(buf.getvalue()).hexdigest()


def count_stats(text: str) -> dict:
    return {
        "lines":      len([l for l in text.split('\n') if l.strip()]),
        "words":      len(text.split()),
        "bold_count": len(re.findall(r'\*\*.+?\*\*', text)),
    }


def convert(image: Image.Image, groq_api_key: str, font_size: int):
    if image is None:
        raise gr.Error("Please upload an image first.")
    if not groq_api_key or not groq_api_key.strip():
        raise gr.Error("Please enter your Groq API key.")

    img_hash = get_image_hash(image)

    yield "🔍 Running Surya OCR (Vision Transformer)…", "", "", None

    if img_hash in _cache:
        raw_text, formatted_text = _cache[img_hash]
        yield "⚡ Loaded from cache — skipping OCR + LLM", raw_text, formatted_text, None
    else:
        try:
            raw_text = run_surya_ocr(image)
        except Exception as e:
            yield f"❌ OCR Error: {type(e).__name__}: {e}", "", "", None
            return

        yield "✨ Formatting with Groq LLaMA 3.3 70B…", raw_text, "", None

        try:
            formatted_text = format_with_groq(raw_text, groq_api_key.strip())
        except Exception as e:
            yield f"❌ Groq Error: {type(e).__name__}: {e}", raw_text, "", None
            return

        _cache[img_hash] = (raw_text, formatted_text)

    yield "📝 Building Word document…", raw_text, formatted_text, None

    docx_bytes = build_docx(formatted_text, font_size)
    import tempfile, os
    tmp_path = os.path.join(tempfile.gettempdir(), "converted.docx")
    with open(tmp_path, "wb") as f:
        f.write(docx_bytes)

    stats = count_stats(formatted_text)
    status = (
        f"✅ Done!  "
        f"{stats['lines']} lines · "
        f"{stats['words']} words · "
        f"{stats['bold_count']} bold segments"
    )
    yield status, raw_text, formatted_text, tmp_path


# ══════════════════════════════════════════════════════════════════════════════
# GRADIO UI
# ══════════════════════════════════════════════════════════════════════════════

CSS = """
body { font-family: 'Segoe UI', sans-serif; }
.title-block { text-align: center; margin-bottom: 8px; }
.title-block h1 { font-size: 2rem; font-weight: 800;
    background: linear-gradient(90deg, #00D4FF, #7C3AED);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent; }
.title-block p  { color: #888; font-size: 0.85rem; letter-spacing: 2px; }
.status-box { border-radius: 8px; padding: 10px 16px;
    font-family: monospace; font-size: 0.85rem; }
footer { display: none !important; }
"""

DESCRIPTION = """
<div class="title-block">
  <h1>📄 Image → Word Converter</h1>
  <p>FAST-NUCES · BSAI · Virtual Company Project · Phase 1 MVP</p>
  <p style="color:#aaa;font-size:0.78rem;">
    OCR: <b>Surya Vision Transformer</b> (local · no quota) &nbsp;|&nbsp;
    LLM: <b>Groq LLaMA 3.3 70B</b> (free · 30 req/min)
  </p>
</div>
"""

with gr.Blocks(css=CSS, title="Image → Word | FAST-NUCES") as demo:

    gr.HTML(DESCRIPTION)

    with gr.Row():
        with gr.Column(scale=1):
            image_input = gr.Image(
                type="pil",
                label="📁 Upload Image (JPG / PNG)",
                height=340,
            )
            groq_key = gr.Textbox(
                label="🔑 Groq API Key",
                placeholder="gsk_...",
                type="password",
                info="Free at console.groq.com — 30 req/min, 14 400 req/day"
            )
            font_size = gr.Slider(
                minimum=10, maximum=18, value=12, step=1,
                label="Font Size (pt)"
            )
            convert_btn = gr.Button("⚡ Convert to Word Document",
                                    variant="primary", size="lg")

        with gr.Column(scale=1):
            status_out = gr.Textbox(
                label="Status", interactive=False,
                elem_classes=["status-box"]
            )
            download_out = gr.File(
                label="⬇ Download Word Document",
                interactive=False,
            )
            with gr.Tabs():
                with gr.Tab("✨ Formatted Text"):
                    formatted_out = gr.Textbox(
                        label="", lines=16, interactive=True,
                        placeholder="Groq-formatted text appears here…"
                    )
                with gr.Tab("🔍 Raw OCR Text"):
                    raw_out = gr.Textbox(
                        label="", lines=16, interactive=False,
                        placeholder="Raw Surya OCR output appears here…"
                    )

    gr.Markdown("""
    ---
    **How it works:**
    `Image` → **Surya OCR** *(ViT, runs locally, zero quota)* → raw text
    → **Groq LLaMA 3.3 70B** *(formats bold / italic / headings)* → **python-docx** → `.docx`
    """)

    convert_btn.click(
        fn=convert,
        inputs=[image_input, groq_key, font_size],
        outputs=[status_out, raw_out, formatted_out, download_out],
    )


if __name__ == "__main__":
    demo.launch()
